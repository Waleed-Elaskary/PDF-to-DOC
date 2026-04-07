"""Core PDF -> DOCX conversion logic.

Engines (pick the best available, or force with --engine):

  word      Uses Microsoft Word (COM automation, Windows only).
            Produces the most editable output: real text, tables, styles.
            Requires Microsoft Word installed + pywin32.
  pdf2docx  Layout-aware Python conversion. Preserves tables/images but places
            content in floating text frames.
  text      Plain text only (paragraph per line, page break per page).

For scanned PDFs, --ocr first produces a searchable PDF via Tesseract, then
runs the chosen engine on it.
"""
from __future__ import annotations

import logging
import os
import sys
import tempfile
from pathlib import Path
from typing import Iterable

import fitz  # PyMuPDF
from docx import Document
from pdf2docx import Converter

logger = logging.getLogger(__name__)

ENGINE_AUTO = "auto"
ENGINE_WORD = "word"
ENGINE_PDF2DOCX = "pdf2docx"
ENGINE_TEXT = "text"
VALID_ENGINES = (ENGINE_AUTO, ENGINE_WORD, ENGINE_PDF2DOCX, ENGINE_TEXT)

_DEFAULT_TESSDATA_CANDIDATES = [
    r"C:\Program Files\Tesseract-OCR\tessdata",
    r"C:\Program Files (x86)\Tesseract-OCR\tessdata",
    "/usr/share/tesseract-ocr/4.00/tessdata",
    "/usr/share/tesseract-ocr/5/tessdata",
    "/usr/local/share/tessdata",
    "/opt/homebrew/share/tessdata",
]


def _find_tessdata() -> str | None:
    env = os.environ.get("TESSDATA_PREFIX")
    if env and Path(env).is_dir():
        return env
    for candidate in _DEFAULT_TESSDATA_CANDIDATES:
        if Path(candidate).is_dir():
            return candidate
    return None


def _word_available() -> bool:
    """True if we can drive MS Word via COM on this machine."""
    if sys.platform != "win32":
        return False
    try:
        import win32com.client  # noqa: F401
        import pythoncom  # noqa: F401
    except ImportError:
        return False
    try:
        import win32com.client as w
        app = w.DispatchEx("Word.Application")
        app.Quit()
        return True
    except Exception as exc:
        logger.debug("Word COM probe failed: %s", exc)
        return False


def resolve_output_path(pdf_path: Path, *, overwrite: bool) -> Path:
    base = pdf_path.with_suffix(".docx")
    if overwrite or not base.exists():
        return base
    counter = 1
    while True:
        candidate = base.with_name(f"{base.stem} ({counter}).docx")
        if not candidate.exists():
            return candidate
        counter += 1


def _pdf_has_text(pdf_path: Path) -> bool:
    try:
        with fitz.open(pdf_path) as pdf:
            for page in pdf:
                if (page.get_text("text") or "").strip():
                    return True
    except Exception as exc:
        logger.warning("Could not probe %s for text: %s", pdf_path.name, exc)
    return False


# --- Engines -----------------------------------------------------------------

def _engine_word(pdf_path: Path, output_path: Path) -> None:
    """Open the PDF in Microsoft Word and save as .docx.

    Word's native PDF importer reflows content into editable text, tables, and
    styles. This is the highest-fidelity editable output on Windows.
    """
    import pythoncom
    import win32com.client as w

    wdFormatDocumentDefault = 16  # .docx

    pythoncom.CoInitialize()
    app = None
    doc = None
    try:
        app = w.DispatchEx("Word.Application")
        app.Visible = False
        app.DisplayAlerts = 0  # wdAlertsNone
        # ConfirmConversions=False suppresses the "convert PDF" dialog.
        doc = app.Documents.Open(
            FileName=str(pdf_path.resolve()),
            ConfirmConversions=False,
            ReadOnly=True,
            AddToRecentFiles=False,
            Visible=False,
        )
        doc.SaveAs(
            FileName=str(output_path.resolve()),
            FileFormat=wdFormatDocumentDefault,
        )
    finally:
        try:
            if doc is not None:
                doc.Close(SaveChanges=0)
        except Exception:
            pass
        try:
            if app is not None:
                app.Quit()
        except Exception:
            pass
        pythoncom.CoUninitialize()


def _engine_pdf2docx(pdf_path: Path, output_path: Path) -> None:
    cv = Converter(str(pdf_path))
    try:
        cv.convert(str(output_path), start=0, end=None)
    finally:
        cv.close()


def _engine_text(pdf_path: Path, output_path: Path) -> None:
    document = Document()
    with fitz.open(pdf_path) as pdf:
        page_count = pdf.page_count
        for index, page in enumerate(pdf):
            text = page.get_text("text") or ""
            if text.strip():
                for line in text.splitlines():
                    document.add_paragraph(line.rstrip())
            else:
                document.add_paragraph("")
            if index < page_count - 1:
                document.add_page_break()
    document.save(output_path)


# --- OCR preprocessing -------------------------------------------------------

def _make_searchable_pdf(
    pdf_path: Path,
    *,
    language: str,
    dpi: int,
    tessdata: str,
) -> Path:
    os.environ["TESSDATA_PREFIX"] = tessdata
    src = fitz.open(pdf_path)
    try:
        out = fitz.open()
        for page in src:
            pix = page.get_pixmap(dpi=dpi)
            ocr_pdf_bytes = pix.pdfocr_tobytes(language=language, tessdata=tessdata)
            ocr_doc = fitz.open("pdf", ocr_pdf_bytes)
            out.insert_pdf(ocr_doc)
            ocr_doc.close()
        tmp_fd, tmp_name = tempfile.mkstemp(suffix=".pdf", prefix="ocr_")
        os.close(tmp_fd)
        out.save(tmp_name)
        out.close()
        return Path(tmp_name)
    finally:
        src.close()


# --- Orchestration -----------------------------------------------------------

def _pick_engine(requested: str) -> str:
    if requested != ENGINE_AUTO:
        return requested
    if _word_available():
        logger.debug("Auto-selected engine: word")
        return ENGINE_WORD
    logger.debug("Auto-selected engine: pdf2docx")
    return ENGINE_PDF2DOCX


def _run_engine(engine: str, pdf_path: Path, output_path: Path) -> None:
    if engine == ENGINE_WORD:
        _engine_word(pdf_path, output_path)
    elif engine == ENGINE_PDF2DOCX:
        _engine_pdf2docx(pdf_path, output_path)
    elif engine == ENGINE_TEXT:
        _engine_text(pdf_path, output_path)
    else:
        raise ValueError(f"Unknown engine: {engine}")


def convert_pdf_to_docx(
    pdf_path: Path,
    *,
    overwrite: bool = False,
    engine: str = ENGINE_AUTO,
    ocr: bool = False,
    ocr_language: str = "eng",
    ocr_dpi: int = 300,
    tessdata: str | None = None,
) -> Path:
    """Convert one PDF into .docx next to the source."""
    pdf_path = Path(pdf_path)
    if not pdf_path.is_file():
        raise FileNotFoundError(f"PDF not found: {pdf_path}")
    if engine not in VALID_ENGINES:
        raise ValueError(
            f"engine must be one of {VALID_ENGINES}, got {engine!r}"
        )

    output_path = resolve_output_path(pdf_path, overwrite=overwrite)
    chosen = _pick_engine(engine)
    logger.info(
        "Converting %s -> %s (engine=%s)",
        pdf_path.name, output_path.name, chosen,
    )

    source_pdf = pdf_path
    tmp_pdf: Path | None = None
    try:
        if not _pdf_has_text(pdf_path):
            if not ocr:
                logger.warning(
                    "%s has no text layer (scanned). Re-run with --ocr.",
                    pdf_path.name,
                )
                Document().save(output_path)
                return output_path

            effective_tessdata = tessdata or _find_tessdata()
            if not effective_tessdata:
                raise RuntimeError(
                    "OCR requested but no tessdata directory found. "
                    "Install Tesseract or set --tessdata / TESSDATA_PREFIX."
                )
            logger.info(
                "No text layer. Running OCR (lang=%s, dpi=%d).",
                ocr_language, ocr_dpi,
            )
            tmp_pdf = _make_searchable_pdf(
                pdf_path,
                language=ocr_language,
                dpi=ocr_dpi,
                tessdata=effective_tessdata,
            )
            source_pdf = tmp_pdf

        _run_engine(chosen, source_pdf, output_path)
        logger.info("Wrote %s", output_path.name)
        return output_path
    finally:
        if tmp_pdf and tmp_pdf.exists():
            try:
                tmp_pdf.unlink()
            except OSError:
                pass


def collect_pdfs(
    inputs: Iterable[Path],
    *,
    recursive: bool = False,
) -> list[Path]:
    """Resolve CLI inputs (files and/or folders) into a list of PDFs.

    - A folder expands to its *.pdf children. Non-recursive by default;
      pass recursive=True to include all subfolders.
    - Files must end with .pdf (case-insensitive).
    """
    resolved: list[Path] = []
    seen: set[Path] = set()
    for item in inputs:
        p = Path(item)
        if p.is_dir():
            iterator = p.rglob("*") if recursive else p.iterdir()
            for child in sorted(iterator):
                if child.is_file() and child.suffix.lower() == ".pdf":
                    rp = child.resolve()
                    if rp not in seen:
                        seen.add(rp)
                        resolved.append(child)
        elif p.is_file():
            if p.suffix.lower() != ".pdf":
                logger.warning("Skipping non-PDF file: %s", p)
                continue
            rp = p.resolve()
            if rp not in seen:
                seen.add(rp)
                resolved.append(p)
        else:
            logger.warning("Input does not exist: %s", p)
    return resolved
